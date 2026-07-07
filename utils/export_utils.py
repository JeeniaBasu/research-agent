"""
=============================================================================
utils/export_utils.py — PDF and DOCX Report Export
=============================================================================
Generates professional PDF and DOCX documents from research reports.
Includes cover pages, tables of contents, structured sections, and references.
=============================================================================
"""

import logging
import os
import re
from datetime import datetime
from io import BytesIO

from config import ExportConfig, AGENT_INSTRUCTIONS

logger = logging.getLogger(__name__)

# Ensure the export directory exists
os.makedirs(ExportConfig.EXPORT_DIR, exist_ok=True)


class ExportService:
    """
    Exports research reports to PDF (ReportLab) and DOCX (python-docx) formats.
    Both methods accept a standardised report dict and return the file path
    or BytesIO object.
    """

    # ------------------------------------------------------------------
    # PDF Export — using ReportLab
    # ------------------------------------------------------------------

    def export_pdf(self, report: dict) -> BytesIO:
        """
        Generate a professional PDF research report.

        Args:
            report: {title, topic, sections: [{heading, content}], references, generated_at}

        Returns:
            BytesIO object containing the PDF data.
        """
        try:
            from reportlab.lib.pagesizes import A4, letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch, cm
            from reportlab.lib import colors
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
                PageBreak, HRFlowable
            )
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=2 * cm,
                leftMargin=2 * cm,
                topMargin=2.5 * cm,
                bottomMargin=2.5 * cm,
                title=report.get("title", "Research Report"),
                author=AGENT_INSTRUCTIONS["name"],
            )

            # Define styles
            styles = getSampleStyleSheet()
            style_title = ParagraphStyle(
                "ReportTitle",
                parent=styles["Title"],
                fontSize=22,
                textColor=colors.HexColor("#1a365d"),
                spaceAfter=12,
                alignment=TA_CENTER,
                fontName="Helvetica-Bold",
            )
            style_subtitle = ParagraphStyle(
                "Subtitle",
                parent=styles["Normal"],
                fontSize=12,
                textColor=colors.HexColor("#4a5568"),
                spaceAfter=6,
                alignment=TA_CENTER,
            )
            style_h1 = ParagraphStyle(
                "H1",
                parent=styles["Heading1"],
                fontSize=16,
                textColor=colors.HexColor("#2d3748"),
                spaceBefore=18,
                spaceAfter=8,
                fontName="Helvetica-Bold",
                borderPad=4,
            )
            style_h2 = ParagraphStyle(
                "H2",
                parent=styles["Heading2"],
                fontSize=13,
                textColor=colors.HexColor("#4a5568"),
                spaceBefore=12,
                spaceAfter=6,
                fontName="Helvetica-Bold",
            )
            style_body = ParagraphStyle(
                "Body",
                parent=styles["Normal"],
                fontSize=10.5,
                textColor=colors.HexColor("#2d3748"),
                spaceAfter=8,
                leading=16,
                alignment=TA_JUSTIFY,
            )
            style_meta = ParagraphStyle(
                "Meta",
                parent=styles["Normal"],
                fontSize=9,
                textColor=colors.HexColor("#718096"),
                spaceAfter=4,
                alignment=TA_CENTER,
            )
            style_ref = ParagraphStyle(
                "Reference",
                parent=styles["Normal"],
                fontSize=9.5,
                textColor=colors.HexColor("#4a5568"),
                spaceAfter=6,
                leftIndent=20,
                firstLineIndent=-20,
                leading=14,
            )

            story = []

            # Cover page
            story.append(Spacer(1, 1.5 * inch))
            story.append(Paragraph(report.get("title", "Research Report"), style_title))
            story.append(Spacer(1, 0.2 * inch))

            if report.get("topic"):
                story.append(Paragraph(f"Topic: {report['topic']}", style_subtitle))

            story.append(Spacer(1, 0.15 * inch))
            generated_at = report.get("generated_at", datetime.now().strftime("%B %d, %Y"))
            story.append(Paragraph(f"Generated by {AGENT_INSTRUCTIONS['name']} | {generated_at}", style_meta))
            story.append(Paragraph("Powered by IBM Granite on IBM Watsonx.ai", style_meta))

            story.append(Spacer(1, 0.3 * inch))
            story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#3b82f6")))
            story.append(PageBreak())

            # Content sections
            for section in report.get("sections", []):
                heading = section.get("heading", "")
                content = section.get("content", "")

                if heading:
                    story.append(Paragraph(heading, style_h1))
                    story.append(HRFlowable(
                        width="100%", thickness=0.5,
                        color=colors.HexColor("#e2e8f0"),
                        spaceAfter=6,
                    ))

                if content:
                    # Parse markdown-style content into paragraphs
                    for block in self._parse_content_blocks(content):
                        if block["type"] == "heading2":
                            story.append(Paragraph(block["text"], style_h2))
                        elif block["type"] == "bullet":
                            story.append(Paragraph(f"• {block['text']}", style_body))
                        elif block["type"] == "numbered":
                            story.append(Paragraph(block["text"], style_body))
                        else:
                            story.append(Paragraph(
                                self._clean_markdown(block["text"]), style_body
                            ))
                        story.append(Spacer(1, 2))

            # References section
            if report.get("references"):
                story.append(PageBreak())
                story.append(Paragraph("References", style_h1))
                story.append(HRFlowable(
                    width="100%", thickness=0.5,
                    color=colors.HexColor("#e2e8f0"),
                    spaceAfter=8,
                ))
                for ref in report["references"]:
                    story.append(Paragraph(ref, style_ref))
                    story.append(Spacer(1, 3))

            # Footer note
            story.append(Spacer(1, 0.3 * inch))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e2e8f0")))
            story.append(Paragraph(
                f"Generated by ResearchMind AI Agent | IBM Watsonx.ai | {generated_at}",
                style_meta,
            ))

            doc.build(story)
            buffer.seek(0)
            return buffer

        except Exception as exc:
            logger.error("PDF export error: %s", exc)
            raise RuntimeError(f"Failed to generate PDF: {exc}") from exc

    # ------------------------------------------------------------------
    # DOCX Export — using python-docx
    # ------------------------------------------------------------------

    def export_docx(self, report: dict) -> BytesIO:
        """
        Generate a professional DOCX research report.

        Args:
            report: {title, topic, sections: [{heading, content}], references, generated_at}

        Returns:
            BytesIO object containing the DOCX data.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            doc = Document()

            # Page margins
            for section in doc.sections:
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)

            # Title
            title_para = doc.add_heading(report.get("title", "Research Report"), 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0] if title_para.runs else title_para.add_run()
            title_run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
            title_run.font.size = Pt(22)

            # Metadata
            if report.get("topic"):
                meta = doc.add_paragraph(f"Topic: {report['topic']}")
                meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                meta.runs[0].font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
                meta.runs[0].font.size = Pt(12)

            generated_at = report.get("generated_at", datetime.now().strftime("%B %d, %Y"))
            meta2 = doc.add_paragraph(
                f"Generated by {AGENT_INSTRUCTIONS['name']} | {generated_at}"
            )
            meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta2.runs[0].font.size = Pt(9)
            meta2.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)

            doc.add_paragraph("Powered by IBM Granite on IBM Watsonx.ai").alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Page break after cover
            doc.add_page_break()

            # Content sections
            for section_data in report.get("sections", []):
                heading = section_data.get("heading", "")
                content = section_data.get("content", "")

                if heading:
                    h = doc.add_heading(heading, level=1)
                    h.runs[0].font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

                if content:
                    for block in self._parse_content_blocks(content):
                        if block["type"] == "heading2":
                            h2 = doc.add_heading(block["text"], level=2)
                            if h2.runs:
                                h2.runs[0].font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
                        elif block["type"] == "bullet":
                            p = doc.add_paragraph(
                                self._clean_markdown(block["text"]),
                                style="List Bullet"
                            )
                        elif block["type"] == "numbered":
                            p = doc.add_paragraph(
                                self._clean_markdown(block["text"]),
                                style="List Number"
                            )
                        elif block["text"].strip():
                            p = doc.add_paragraph(self._clean_markdown(block["text"]))
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            for run in p.runs:
                                run.font.size = Pt(11)

            # References
            if report.get("references"):
                doc.add_page_break()
                doc.add_heading("References", level=1)
                for ref in report["references"]:
                    p = doc.add_paragraph(ref, style="List Number")
                    for run in p.runs:
                        run.font.size = Pt(10)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        except Exception as exc:
            logger.error("DOCX export error: %s", exc)
            raise RuntimeError(f"Failed to generate DOCX: {exc}") from exc

    # ------------------------------------------------------------------
    # Markdown parsing helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_content_blocks(text: str) -> list:
        """
        Convert markdown-style text into a list of typed blocks
        for structured document generation.
        """
        blocks = []
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith("## "):
                blocks.append({"type": "heading2", "text": stripped[3:].strip()})
            elif stripped.startswith("### "):
                blocks.append({"type": "heading2", "text": stripped[4:].strip()})
            elif stripped.startswith(("- ", "* ", "• ")):
                blocks.append({"type": "bullet", "text": stripped[2:].strip()})
            elif re.match(r"^\d+\.\s", stripped):
                blocks.append({"type": "numbered", "text": re.sub(r"^\d+\.\s", "", stripped)})
            else:
                blocks.append({"type": "paragraph", "text": stripped})

        return blocks

    @staticmethod
    def _clean_markdown(text: str) -> str:
        """Strip markdown syntax from text for plain document rendering."""
        # Remove bold/italic markers
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"__(.+?)__", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        # Remove markdown links
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
        return text.strip()
